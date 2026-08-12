import io
import zipfile

import pytest
from docx import Document
from pypdf import PdfWriter

from farmers_chatbot import documents


def _zip_bytes(entries: dict[str, bytes]) -> bytes:
    output = io.BytesIO()
    with zipfile.ZipFile(output, "w", compression=zipfile.ZIP_DEFLATED) as archive:
        for name, content in entries.items():
            archive.writestr(name, content)
    return output.getvalue()


def _encrypted_pdf() -> bytes:
    output = io.BytesIO()
    writer = PdfWriter()
    writer.add_blank_page(width=100, height=100)
    writer.encrypt("secret")
    writer.write(output)
    return output.getvalue()


def test_valid_docx_is_extracted_through_isolated_worker() -> None:
    output = io.BytesIO()
    document = Document()
    document.add_paragraph("Soil pH is 6.2.")
    document.add_paragraph("درجة حموضة التربة 6.2")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Irrigation"
    table.cell(0, 1).text = "Weekly"
    document.save(output)

    text = documents.extract_document_text("field-notes.docx", output.getvalue())

    assert "Soil pH is 6.2." in text
    assert "درجة حموضة التربة 6.2" in text
    assert "Irrigation | Weekly" in text


def test_malformed_pdf_fails_with_sanitized_error() -> None:
    with pytest.raises(ValueError, match="PDF document could not be read"):
        documents.extract_document_text("malformed.pdf", b"%PDF-not-a-pdf")


def test_encrypted_pdf_is_rejected() -> None:
    with pytest.raises(ValueError, match="Encrypted PDFs are not accepted"):
        documents.extract_document_text("encrypted.pdf", _encrypted_pdf())


def test_pdf_parser_timeout_kills_worker(monkeypatch: pytest.MonkeyPatch) -> None:
    monkeypatch.setattr(documents, "DOCUMENT_PARSE_TIMEOUT_SECONDS", 1e-9)

    with pytest.raises(ValueError, match="exceeded the safety time limit"):
        documents.extract_document_text("slow.pdf", b"%PDF-not-a-pdf")


def test_oversized_pdf_is_rejected_before_parser_spawn(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    monkeypatch.setattr(documents, "MAX_PROJECT_FILE_BYTES", 16)

    with pytest.raises(ValueError, match="no larger than 10 MB"):
        documents.extract_document_text("large.pdf", b"%PDF-" + (b"x" * 12))


@pytest.mark.parametrize(
    "unsafe_name",
    ["../escape.xml", "folder/../../escape.xml", "C:\\escape.xml"],
)
def test_office_archive_rejects_unsafe_paths(unsafe_name: str) -> None:
    archive = _zip_bytes({"[Content_Types].xml": b"types", unsafe_name: b"bad"})

    with pytest.raises(ValueError, match="unsafe path"):
        documents._validate_zip(archive)


def test_office_archive_rejects_suspicious_compression_ratio() -> None:
    archive = _zip_bytes({"word/document.xml": b"A" * 200_000})

    with pytest.raises(ValueError, match="suspicious compression ratio"):
        documents._validate_zip(archive)


def test_office_archive_rejects_symbolic_links() -> None:
    output = io.BytesIO()
    link = zipfile.ZipInfo("word/linked.xml")
    link.create_system = 3
    link.external_attr = 0o120777 << 16
    with zipfile.ZipFile(output, "w") as archive:
        archive.writestr(link, "../../outside")

    with pytest.raises(ValueError, match="unsafe symbolic link"):
        documents._validate_zip(output.getvalue())


def test_malformed_office_archive_is_rejected() -> None:
    with pytest.raises(ValueError, match="not a valid ZIP-based file"):
        documents.extract_document_text("malformed.docx", b"not-a-zip")

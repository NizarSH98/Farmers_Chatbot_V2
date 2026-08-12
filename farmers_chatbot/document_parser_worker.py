"""Resource-limited worker for parsing untrusted project documents.

Launched as a standalone child process by ``documents.py``. Keeping this
module free of application imports avoids loading credentials or service state
into the parser process.
"""

from __future__ import annotations

import io
import json
import os
import re
import sys
from collections.abc import Iterable


def _apply_posix_limits(memory_bytes: int, cpu_seconds: int) -> None:
    import resource

    resource.setrlimit(resource.RLIMIT_AS, (memory_bytes, memory_bytes))
    resource.setrlimit(resource.RLIMIT_CPU, (cpu_seconds, cpu_seconds))
    resource.setrlimit(resource.RLIMIT_CORE, (0, 0))
    output_bytes = 2 * 1024 * 1024
    resource.setrlimit(resource.RLIMIT_FSIZE, (output_bytes, output_bytes))


def _apply_windows_limits(memory_bytes: int, cpu_seconds: int) -> None:
    import ctypes
    from ctypes import wintypes

    class IoCounters(ctypes.Structure):
        _fields_ = [
            ("ReadOperationCount", ctypes.c_ulonglong),
            ("WriteOperationCount", ctypes.c_ulonglong),
            ("OtherOperationCount", ctypes.c_ulonglong),
            ("ReadTransferCount", ctypes.c_ulonglong),
            ("WriteTransferCount", ctypes.c_ulonglong),
            ("OtherTransferCount", ctypes.c_ulonglong),
        ]

    class BasicLimitInformation(ctypes.Structure):
        _fields_ = [
            ("PerProcessUserTimeLimit", ctypes.c_longlong),
            ("PerJobUserTimeLimit", ctypes.c_longlong),
            ("LimitFlags", wintypes.DWORD),
            ("MinimumWorkingSetSize", ctypes.c_size_t),
            ("MaximumWorkingSetSize", ctypes.c_size_t),
            ("ActiveProcessLimit", wintypes.DWORD),
            ("Affinity", ctypes.c_size_t),
            ("PriorityClass", wintypes.DWORD),
            ("SchedulingClass", wintypes.DWORD),
        ]

    class ExtendedLimitInformation(ctypes.Structure):
        _fields_ = [
            ("BasicLimitInformation", BasicLimitInformation),
            ("IoInfo", IoCounters),
            ("ProcessMemoryLimit", ctypes.c_size_t),
            ("JobMemoryLimit", ctypes.c_size_t),
            ("PeakProcessMemoryUsed", ctypes.c_size_t),
            ("PeakJobMemoryUsed", ctypes.c_size_t),
        ]

    kernel32 = ctypes.WinDLL("kernel32", use_last_error=True)
    kernel32.CreateJobObjectW.argtypes = [ctypes.c_void_p, wintypes.LPCWSTR]
    kernel32.CreateJobObjectW.restype = wintypes.HANDLE
    kernel32.SetInformationJobObject.argtypes = [
        wintypes.HANDLE,
        ctypes.c_int,
        ctypes.c_void_p,
        wintypes.DWORD,
    ]
    kernel32.SetInformationJobObject.restype = wintypes.BOOL
    kernel32.AssignProcessToJobObject.argtypes = [wintypes.HANDLE, wintypes.HANDLE]
    kernel32.AssignProcessToJobObject.restype = wintypes.BOOL
    kernel32.GetCurrentProcess.restype = wintypes.HANDLE

    job = kernel32.CreateJobObjectW(None, None)
    if not job:
        raise OSError(ctypes.get_last_error(), "Unable to create parser job object")
    limits = ExtendedLimitInformation()
    limits.BasicLimitInformation.PerProcessUserTimeLimit = cpu_seconds * 10_000_000
    limits.BasicLimitInformation.LimitFlags = 0x2 | 0x100 | 0x2000
    limits.ProcessMemoryLimit = memory_bytes
    if not kernel32.SetInformationJobObject(
        job,
        9,
        ctypes.byref(limits),
        ctypes.sizeof(limits),
    ):
        raise OSError(ctypes.get_last_error(), "Unable to configure parser job object")
    if not kernel32.AssignProcessToJobObject(job, kernel32.GetCurrentProcess()):
        raise OSError(ctypes.get_last_error(), "Unable to isolate parser process")
    globals()["_WINDOWS_JOB_HANDLE"] = job


def _apply_resource_limits(memory_bytes: int, cpu_seconds: int) -> None:
    if os.name == "nt":
        _apply_windows_limits(memory_bytes, cpu_seconds)
    else:
        _apply_posix_limits(memory_bytes, cpu_seconds)


def _bounded_join(parts: Iterable[str], separator: str, max_characters: int) -> str:
    output: list[str] = []
    used = 0
    for part in parts:
        if not part:
            continue
        available = max_characters - used
        if output:
            if available <= len(separator):
                break
            output.append(separator)
            used += len(separator)
            available -= len(separator)
        output.append(part[:available])
        used += min(len(part), available)
        if used >= max_characters:
            break
    return "".join(output)


def _clean_text(text: str, max_characters: int) -> str:
    cleaned = re.sub(r"\x00", "", text)
    cleaned = re.sub(r"[ \t]+", " ", cleaned)
    cleaned = re.sub(r"\n{3,}", "\n\n", cleaned).strip()
    if not cleaned:
        raise ValueError("No readable text was found in the document")
    return cleaned[:max_characters]


def _extract_pdf(data: bytes, max_characters: int, max_pages: int) -> str:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(data), strict=True)
    if reader.is_encrypted:
        raise ValueError("Encrypted PDFs are not accepted")

    def pages() -> Iterable[str]:
        for page_number, page in enumerate(reader.pages):
            if page_number >= max_pages:
                break
            yield page.extract_text() or ""

    return _bounded_join(pages(), "\n\n", max_characters)


def _extract_docx(data: bytes, max_characters: int) -> str:
    from docx import Document

    document = Document(io.BytesIO(data))

    def parts() -> Iterable[str]:
        yield from (paragraph.text for paragraph in document.paragraphs)
        for table in document.tables:
            for row in table.rows:
                yield " | ".join(cell.text for cell in row.cells)

    return _bounded_join(parts(), "\n", max_characters)


def _extract_xlsx(data: bytes, max_characters: int) -> str:
    from openpyxl import load_workbook

    workbook = load_workbook(io.BytesIO(data), read_only=True, data_only=True)

    def lines() -> Iterable[str]:
        for sheet in workbook.worksheets[:20]:
            yield f"Sheet: {sheet.title}"
            for row in sheet.iter_rows(max_row=5000, values_only=True):
                values = [str(value) for value in row if value not in {None, ""}]
                if values:
                    yield " | ".join(values)

    try:
        return _bounded_join(lines(), "\n", max_characters)
    finally:
        workbook.close()


def _extract(extension: str, data: bytes, max_characters: int, max_pages: int) -> str:
    if extension == ".pdf":
        text = _extract_pdf(data, max_characters, max_pages)
    elif extension == ".docx":
        text = _extract_docx(data, max_characters)
    elif extension == ".xlsx":
        text = _extract_xlsx(data, max_characters)
    else:
        raise ValueError("Unsupported isolated document type")
    return _clean_text(text, max_characters)


def _safe_error(extension: str, exc: Exception) -> str:
    message = str(exc)
    if isinstance(exc, ValueError) and message in {
        "Encrypted PDFs are not accepted",
        "No readable text was found in the document",
    }:
        return message
    labels = {
        ".pdf": "PDF document could not be read",
        ".docx": "Word document could not be read",
        ".xlsx": "Spreadsheet could not be read",
    }
    return labels.get(extension, "Document could not be read")


def main() -> int:
    if len(sys.argv) != 7:
        return 2
    extension = sys.argv[1]
    try:
        max_characters = int(sys.argv[2])
        max_pages = int(sys.argv[3])
        max_input_bytes = int(sys.argv[4])
        memory_bytes = int(sys.argv[5])
        cpu_seconds = int(sys.argv[6])
        _apply_resource_limits(memory_bytes, cpu_seconds)
        data = sys.stdin.buffer.read(max_input_bytes + 1)
        if not data or len(data) > max_input_bytes:
            raise ValueError("Document exceeds the parser input limit")
        text = _extract(extension, data, max_characters, max_pages)
        response = {"ok": True, "text": text}
        status = 0
    except Exception as exc:  # noqa: BLE001 - process boundary sanitizes all failures
        response = {"ok": False, "error": _safe_error(extension, exc)}
        status = 2
    encoded_response = json.dumps(response, ensure_ascii=False).encode("utf-8")
    sys.stdout.buffer.write(encoded_response)
    return status


if __name__ == "__main__":
    raise SystemExit(main())

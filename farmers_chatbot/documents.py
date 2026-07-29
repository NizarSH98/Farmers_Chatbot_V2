"""Validated project-document ingestion and scoped retrieval."""

from __future__ import annotations

import csv
import hashlib
import io
import mimetypes
import re
import uuid
import zipfile
from dataclasses import dataclass
from pathlib import Path
from typing import Any

import numpy as np
from docx import Document
from openpyxl import load_workbook
from PyPDF2 import PdfReader
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.pipeline import FeatureUnion
from sklearn.preprocessing import normalize

from .config import MAX_PROJECT_FILE_BYTES, MAX_PROJECT_FILES
from .pilot_store import PilotStore
from .storage_backends import PrivateFileStorage

ALLOWED_EXTENSIONS = {".pdf", ".docx", ".txt", ".csv", ".xlsx"}
ALLOWED_MIME_TYPES = {
    ".pdf": {"application/pdf"},
    ".docx": {
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    },
    ".txt": {"text/plain", "application/octet-stream"},
    ".csv": {"text/csv", "application/vnd.ms-excel", "text/plain"},
    ".xlsx": {
        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    },
}
MAX_EXTRACTED_CHARACTERS = 250_000
MAX_ARCHIVE_FILES = 1000
MAX_ARCHIVE_UNCOMPRESSED = 50 * 1024 * 1024


@dataclass(frozen=True)
class ProjectSearchResult:
    chunk_id: str
    document_id: str
    filename: str
    text: str
    score: float
    source_type: str = "user_project_document"

    def to_dict(self) -> dict[str, Any]:
        return {
            "chunk_id": self.chunk_id,
            "document_id": self.document_id,
            "filename": self.filename,
            "text": self.text,
            "score": self.score,
            "source_type": self.source_type,
        }


def _clean_filename(filename: str) -> str:
    name = Path(filename).name
    cleaned = re.sub(r"[^A-Za-z0-9._\-\u0600-\u06ff ]+", "_", name).strip()
    return cleaned[:180] or "document"


def _validate_zip(data: bytes) -> None:
    try:
        with zipfile.ZipFile(io.BytesIO(data)) as archive:
            infos = archive.infolist()
            if len(infos) > MAX_ARCHIVE_FILES:
                raise ValueError("Office file contains too many embedded files")
            total = sum(info.file_size for info in infos)
            if total > MAX_ARCHIVE_UNCOMPRESSED:
                raise ValueError("Office file expands beyond the pilot safety limit")
            for info in infos:
                candidate = Path(info.filename)
                if candidate.is_absolute() or ".." in candidate.parts:
                    raise ValueError("Office file contains an unsafe path")
    except zipfile.BadZipFile as exc:
        raise ValueError("Office document is not a valid ZIP-based file") from exc


def _decode_text(data: bytes) -> str:
    for encoding in ("utf-8-sig", "utf-8", "cp1256", "latin-1"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    raise ValueError("Text encoding is not supported")


def extract_document_text(filename: str, data: bytes) -> str:
    extension = Path(filename).suffix.lower()
    if extension == ".pdf":
        if not data.startswith(b"%PDF-"):
            raise ValueError("PDF signature is invalid")
        reader = PdfReader(io.BytesIO(data))
        if reader.is_encrypted:
            raise ValueError("Encrypted PDFs are not accepted")
        text = "\n\n".join((page.extract_text() or "") for page in reader.pages[:200])
    elif extension == ".docx":
        _validate_zip(data)
        document = Document(io.BytesIO(data))
        parts = [paragraph.text for paragraph in document.paragraphs]
        for table in document.tables:
            parts.extend(
                " | ".join(cell.text for cell in row.cells) for row in table.rows
            )
        text = "\n".join(parts)
    elif extension == ".xlsx":
        _validate_zip(data)
        workbook = load_workbook(
            io.BytesIO(data),
            read_only=True,
            data_only=True,
        )
        lines: list[str] = []
        for sheet in workbook.worksheets[:20]:
            lines.append(f"Sheet: {sheet.title}")
            for row in sheet.iter_rows(max_row=5000, values_only=True):
                values = [str(value) for value in row if value not in {None, ""}]
                if values:
                    lines.append(" | ".join(values))
        workbook.close()
        text = "\n".join(lines)
    elif extension == ".csv":
        decoded = _decode_text(data)
        rows = csv.reader(io.StringIO(decoded))
        text = "\n".join(" | ".join(row) for _, row in zip(range(5001), rows))
    elif extension == ".txt":
        text = _decode_text(data)
    else:
        raise ValueError("Unsupported document type")

    cleaned = re.sub(r"\x00", "", text)
    cleaned = re.sub(r"[ \t]+", " ", cleaned)
    cleaned = re.sub(r"\n{3,}", "\n\n", cleaned).strip()
    if not cleaned:
        raise ValueError("No readable text was found in the document")
    return cleaned[:MAX_EXTRACTED_CHARACTERS]


def chunk_text(text: str, size: int = 1400, overlap: int = 180) -> list[str]:
    paragraphs = [part.strip() for part in re.split(r"\n\s*\n", text) if part.strip()]
    chunks: list[str] = []
    current = ""
    for paragraph in paragraphs:
        if len(current) + len(paragraph) + 2 <= size:
            current = f"{current}\n\n{paragraph}".strip()
            continue
        if current:
            chunks.append(current)
        if len(paragraph) <= size:
            current = paragraph
            continue
        start = 0
        while start < len(paragraph):
            end = min(len(paragraph), start + size)
            chunks.append(paragraph[start:end])
            if end == len(paragraph):
                break
            start = max(start + 1, end - overlap)
        current = ""
    if current:
        chunks.append(current)
    return chunks[:250]


class DocumentService:
    def __init__(self, store: PilotStore, storage: PrivateFileStorage) -> None:
        self.store = store
        self.storage = storage

    def ingest(
        self,
        owner_user_id: str,
        project_id: str,
        *,
        filename: str,
        data: bytes,
        mime_type: str | None = None,
    ) -> str:
        if self.store.document_count(owner_user_id, project_id) >= MAX_PROJECT_FILES:
            raise ValueError(f"A project can contain at most {MAX_PROJECT_FILES} files")
        if not data or len(data) > MAX_PROJECT_FILE_BYTES:
            raise ValueError("File must be non-empty and no larger than 10 MB")

        cleaned_name = _clean_filename(filename)
        extension = Path(cleaned_name).suffix.lower()
        if extension not in ALLOWED_EXTENSIONS:
            raise ValueError("Allowed files: PDF, DOCX, TXT, CSV, and XLSX")
        supplied_mime = (mime_type or "").split(";", 1)[0].lower()
        if supplied_mime and supplied_mime not in ALLOWED_MIME_TYPES[extension]:
            raise ValueError("File extension and media type do not match")

        text = extract_document_text(cleaned_name, data)
        chunks = chunk_text(text)
        digest = hashlib.sha256(data).hexdigest()
        storage_path = (
            f"users/{owner_user_id}/projects/{project_id}/documents/"
            f"{uuid.uuid4()}-{cleaned_name}"
        )
        detected_mime = (
            supplied_mime
            or mimetypes.guess_type(cleaned_name)[0]
            or "application/octet-stream"
        )
        self.storage.put(storage_path, data, detected_mime)
        try:
            return self.store.add_document(
                owner_user_id,
                project_id,
                filename=cleaned_name,
                mime_type=detected_mime,
                storage_path=storage_path,
                sha256=digest,
                size_bytes=len(data),
                chunks=chunks,
            )
        except Exception:
            self.storage.delete(storage_path)
            raise

    def delete(
        self,
        owner_user_id: str,
        project_id: str,
        document_id: str,
    ) -> None:
        path = self.store.delete_document(owner_user_id, project_id, document_id)
        self.storage.delete(path)


def search_project_chunks(
    chunks: list[dict[str, Any]],
    query: str,
    *,
    top_k: int = 5,
    min_score: float = 0.01,
) -> list[ProjectSearchResult]:
    if not chunks or not query.strip():
        return []
    vectorizer = FeatureUnion(
        [
            (
                "words",
                TfidfVectorizer(
                    ngram_range=(1, 2),
                    sublinear_tf=True,
                    token_pattern=r"(?u)\b\w[\w'-]+\b",
                ),
            ),
            (
                "characters",
                TfidfVectorizer(
                    analyzer="char_wb",
                    ngram_range=(3, 5),
                    sublinear_tf=True,
                    max_features=20000,
                ),
            ),
        ]
    )
    texts = [f"{chunk['filename']} {chunk['text_content']}" for chunk in chunks]
    try:
        vectors = normalize(vectorizer.fit_transform(texts))
        query_vector = normalize(vectorizer.transform([query]))
    except ValueError:
        return []
    scores = (query_vector @ vectors.T).toarray().ravel()
    ranked = np.argsort(scores)[::-1]
    results = []
    for index in ranked[: max(1, min(int(top_k), 10))]:
        score = float(scores[int(index)])
        if score < min_score:
            continue
        chunk = chunks[int(index)]
        results.append(
            ProjectSearchResult(
                chunk_id=str(chunk["id"]),
                document_id=str(chunk["document_id"]),
                filename=str(chunk["filename"]),
                text=str(chunk["text_content"]),
                score=score,
            )
        )
    return results

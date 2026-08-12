"""Deterministic, bounded farmer artifact builders."""

from __future__ import annotations

import io
import math
import re
import uuid
from dataclasses import dataclass
from datetime import UTC, date, datetime
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from openpyxl import Workbook
from openpyxl.styles import Alignment, Font, PatternFill

from .config import MAX_ARTIFACTS_PER_USER_DAY
from .pilot_store import PilotStore
from .storage_backends import PrivateFileStorage

DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
XLSX_MIME = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"


@dataclass(frozen=True)
class ArtifactReference:
    artifact_id: str
    artifact_type: str
    filename: str
    mime_type: str

    def to_dict(self) -> dict[str, str]:
        return {
            "artifact_id": self.artifact_id,
            "artifact_type": self.artifact_type,
            "filename": self.filename,
            "mime_type": self.mime_type,
        }


def _safe_filename(value: str, extension: str) -> str:
    base = re.sub(r"[^A-Za-z0-9_\-\u0600-\u06ff ]+", "_", value).strip()
    base = re.sub(r"\s+", "_", base)[:80] or "RAISE_artifact"
    return f"{base}{extension}"


def _safe_cell(value: Any) -> Any:
    if isinstance(value, str) and value.startswith(("=", "+", "-", "@")):
        return f"'{value}"
    return value


def _limited_text(value: str, maximum: int = 5000) -> str:
    cleaned = " ".join((value or "").split()).strip()
    if not cleaned:
        raise ValueError("Artifact text cannot be empty")
    return cleaned[:maximum]


def _limited_items(items: list[Any], maximum: int = 50) -> list[Any]:
    if not items:
        raise ValueError("Artifact requires at least one item")
    return items[:maximum]


def _evidence_ids(values: list[str] | None, maximum: int = 50) -> list[str]:
    result: list[str] = []
    for raw_value in (values or [])[:maximum]:
        value = str(raw_value).strip()
        if not re.fullmatch(r"[A-Za-z0-9][A-Za-z0-9._:-]{0,127}", value):
            raise ValueError(f"Invalid evidence ID: {value[:40]}")
        if value not in result:
            result.append(value)
    return result


def _nonnegative_number(value: Any, label: str) -> float:
    try:
        number = float(value)
    except (TypeError, ValueError) as exc:
        raise ValueError(f"{label} must be a number") from exc
    if not math.isfinite(number) or number < 0:
        raise ValueError(f"{label} must be finite and non-negative")
    return number


def _optional_iso_date(value: str | None) -> str | None:
    if value is None or not value.strip():
        return None
    normalized = value.strip()
    try:
        return date.fromisoformat(normalized).isoformat()
    except ValueError as exc:
        raise ValueError("as_of_date must use YYYY-MM-DD") from exc


def _provenance_metadata(
    evidence_ids: list[str],
    legacy_sources: list[str],
) -> dict[str, Any]:
    if evidence_ids:
        status = "evidence_linked"
    elif legacy_sources:
        status = "legacy_unstructured"
    else:
        status = "unlinked"
    return {
        "evidence_ids": evidence_ids,
        "legacy_sources": legacy_sources[:30],
        "provenance_status": status,
    }


def _set_paragraph_direction(paragraph: Any, arabic: bool) -> None:
    if not arabic:
        return
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    properties = paragraph._p.get_or_add_pPr()
    bidi = OxmlElement("w:bidi")
    bidi.set(qn("w:val"), "1")
    properties.append(bidi)


def _docx_bytes(
    *,
    title: str,
    sections: list[tuple[str, str | list[str]]],
    language: str,
    sources: list[str],
    evidence_ids: list[str],
    assumptions: list[str],
) -> bytes:
    arabic = language == "arabic"
    document = Document()
    heading = document.add_heading(_limited_text(title, 200), 0)
    _set_paragraph_direction(heading, arabic)
    generated = datetime.now(UTC).date().isoformat()
    paragraph = document.add_paragraph(f"Generated / أُنشئ: {generated}")
    _set_paragraph_direction(paragraph, arabic)

    for heading_text, content in sections:
        section_heading = document.add_heading(_limited_text(heading_text, 200), 1)
        _set_paragraph_direction(section_heading, arabic)
        if isinstance(content, list):
            for item in _limited_items(content):
                item_paragraph = document.add_paragraph(
                    _limited_text(str(item)),
                    style="List Bullet",
                )
                _set_paragraph_direction(item_paragraph, arabic)
        else:
            item_paragraph = document.add_paragraph(_limited_text(content))
            _set_paragraph_direction(item_paragraph, arabic)

    if assumptions:
        assumptions_heading = document.add_heading("Assumptions / الافتراضات", 1)
        _set_paragraph_direction(assumptions_heading, arabic)
        for item in assumptions[:20]:
            item_paragraph = document.add_paragraph(
                _limited_text(str(item)),
                style="List Bullet",
            )
            _set_paragraph_direction(item_paragraph, arabic)
    if sources:
        sources_heading = document.add_heading("Sources / المصادر", 1)
        _set_paragraph_direction(sources_heading, arabic)
        for item in sources[:20]:
            item_paragraph = document.add_paragraph(
                _limited_text(str(item)),
                style="List Bullet",
            )
            _set_paragraph_direction(item_paragraph, arabic)

    if evidence_ids:
        evidence_heading = document.add_heading("Evidence IDs", 1)
        _set_paragraph_direction(evidence_heading, arabic)
        for item in evidence_ids[:50]:
            item_paragraph = document.add_paragraph(item, style="List Bullet")
            _set_paragraph_direction(item_paragraph, arabic)

    warning = document.add_paragraph(
        "Pilot decision-support output. Verify high-risk agronomic, pesticide, "
        "veterinary, food-safety, financial, and regulatory decisions with a "
        "qualified local professional."
    )
    warning.runs[0].italic = True
    _set_paragraph_direction(warning, arabic)
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


class ArtifactService:
    def __init__(
        self,
        store: PilotStore,
        storage: PrivateFileStorage,
        *,
        owner_user_id: str,
        project_id: str | None = None,
        conversation_id: str | None = None,
    ) -> None:
        self.store = store
        self.storage = storage
        self.owner_user_id = owner_user_id
        self.project_id = project_id
        self.conversation_id = conversation_id

    def _save(
        self,
        artifact_type: str,
        filename: str,
        mime_type: str,
        data: bytes,
        metadata: dict[str, Any],
    ) -> ArtifactReference:
        if self.store.artifacts_today(self.owner_user_id) >= MAX_ARTIFACTS_PER_USER_DAY:
            raise ValueError("Daily artifact generation limit reached")
        storage_path = f"users/{self.owner_user_id}/artifacts/{uuid.uuid4()}-{filename}"
        self.storage.put(storage_path, data, mime_type)
        try:
            artifact_id = self.store.add_artifact(
                self.owner_user_id,
                artifact_type=artifact_type,
                filename=filename,
                mime_type=mime_type,
                storage_path=storage_path,
                project_id=self.project_id,
                conversation_id=self.conversation_id,
                metadata=metadata,
            )
        except Exception:
            self.storage.delete(storage_path)
            raise
        return ArtifactReference(artifact_id, artifact_type, filename, mime_type)

    def generate_farm_action_plan(
        self,
        *,
        title: str,
        context: str,
        actions: list[str],
        assumptions: list[str] | None = None,
        sources: list[str] | None = None,
        evidence_ids: list[str] | None = None,
        language: str = "english",
    ) -> dict[str, Any]:
        linked_evidence = _evidence_ids(evidence_ids)
        legacy_sources = sources or []
        data = _docx_bytes(
            title=title,
            sections=[
                ("Context / السياق", context),
                ("Action plan / خطة العمل", actions),
            ],
            language=language,
            sources=legacy_sources,
            evidence_ids=linked_evidence,
            assumptions=assumptions or [],
        )
        filename = _safe_filename(title, ".docx")
        return self._save(
            "farm_action_plan",
            filename,
            DOCX_MIME,
            data,
            {
                "language": language,
                "item_count": len(actions),
                **_provenance_metadata(linked_evidence, legacy_sources),
            },
        ).to_dict()

    def generate_inspection_checklist(
        self,
        *,
        title: str,
        context: str,
        checks: list[str],
        escalation_signs: list[str] | None = None,
        sources: list[str] | None = None,
        evidence_ids: list[str] | None = None,
        language: str = "english",
    ) -> dict[str, Any]:
        linked_evidence = _evidence_ids(evidence_ids)
        legacy_sources = sources or []
        sections: list[tuple[str, str | list[str]]] = [
            ("Context / السياق", context),
            ("Inspection checks / نقاط الفحص", [f"☐ {item}" for item in checks]),
        ]
        if escalation_signs:
            sections.append(("Escalate when / متى تجب الإحالة", escalation_signs))
        data = _docx_bytes(
            title=title,
            sections=sections,
            language=language,
            sources=legacy_sources,
            evidence_ids=linked_evidence,
            assumptions=[],
        )
        filename = _safe_filename(title, ".docx")
        return self._save(
            "inspection_checklist",
            filename,
            DOCX_MIME,
            data,
            {
                "language": language,
                "item_count": len(checks),
                **_provenance_metadata(linked_evidence, legacy_sources),
            },
        ).to_dict()

    def generate_expert_referral_brief(
        self,
        *,
        title: str,
        situation: str,
        observations: list[str],
        question_for_expert: str,
        urgent_signs: list[str] | None = None,
        sources: list[str] | None = None,
        evidence_ids: list[str] | None = None,
        language: str = "english",
    ) -> dict[str, Any]:
        linked_evidence = _evidence_ids(evidence_ids)
        legacy_sources = sources or []
        sections: list[tuple[str, str | list[str]]] = [
            ("Situation / الحالة", situation),
            ("Observed information / المعلومات المرصودة", observations),
            ("Decision needed / القرار المطلوب", question_for_expert),
        ]
        if urgent_signs:
            sections.append(("Urgent signs / مؤشرات عاجلة", urgent_signs))
        data = _docx_bytes(
            title=title,
            sections=sections,
            language=language,
            sources=legacy_sources,
            evidence_ids=linked_evidence,
            assumptions=[],
        )
        filename = _safe_filename(title, ".docx")
        return self._save(
            "expert_referral_brief",
            filename,
            DOCX_MIME,
            data,
            {
                "language": language,
                **_provenance_metadata(linked_evidence, legacy_sources),
            },
        ).to_dict()

    def generate_crop_calendar(
        self,
        *,
        title: str,
        entries: list[dict[str, Any]],
        assumptions: list[str] | None = None,
        sources: list[str] | None = None,
        evidence_ids: list[str] | None = None,
        language: str = "english",
    ) -> dict[str, Any]:
        linked_evidence = _evidence_ids(evidence_ids)
        legacy_sources = sources or []
        rows = _limited_items(entries, 60)
        workbook = Workbook()
        sheet = workbook.active
        sheet.title = "Crop calendar"
        headers = ["Period", "Activity", "Decision trigger", "Risk / note", "Source"]
        sheet.append(headers)
        for cell in sheet[1]:
            cell.font = Font(bold=True, color="FFFFFF")
            cell.fill = PatternFill("solid", fgColor="467A52")
            cell.alignment = Alignment(wrap_text=True)
        for entry in rows:
            sheet.append(
                [
                    _safe_cell(str(entry.get("period") or "")),
                    _safe_cell(str(entry.get("activity") or "")),
                    _safe_cell(str(entry.get("trigger") or "")),
                    _safe_cell(str(entry.get("risk") or "")),
                    _safe_cell(str(entry.get("source") or "")),
                ]
            )
        for column in "ABCDE":
            sheet.column_dimensions[column].width = 28
        meta = workbook.create_sheet("Assumptions and sources")
        meta.append(["Generated", datetime.now(UTC).isoformat()])
        meta.append(["Language", language])
        meta.append(["Assumptions"])
        for item in (assumptions or [])[:20]:
            meta.append([_safe_cell(item)])
        meta.append(["Sources"])
        for item in legacy_sources[:20]:
            meta.append([_safe_cell(item)])
        meta.append(["Evidence IDs"])
        for item in linked_evidence:
            meta.append([item])
        buffer = io.BytesIO()
        workbook.save(buffer)
        filename = _safe_filename(title, ".xlsx")
        return self._save(
            "crop_calendar",
            filename,
            XLSX_MIME,
            buffer.getvalue(),
            {
                "language": language,
                "entry_count": len(rows),
                **_provenance_metadata(linked_evidence, legacy_sources),
            },
        ).to_dict()

    def calculate_enterprise_budget(
        self,
        *,
        title: str,
        currency: str,
        costs: list[dict[str, Any]],
        revenues: list[dict[str, Any]],
        assumptions: list[str] | None = None,
        sources: list[str] | None = None,
        evidence_ids: list[str] | None = None,
        geography: str | None = None,
        as_of_date: str | None = None,
        financing_cost: float = 0,
        depreciation_cost: float = 0,
        sensitivity_scenarios: list[dict[str, Any]] | None = None,
        impacts: list[dict[str, Any]] | None = None,
    ) -> dict[str, Any]:
        currency_value = _limited_text(currency, 20)
        geography_value = (
            _limited_text(geography, 200) if geography and geography.strip() else None
        )
        effective_date = _optional_iso_date(as_of_date)
        linked_evidence = _evidence_ids(evidence_ids)
        legacy_sources = sources or []
        assumption_rows = [
            _limited_text(str(item), 1000) for item in (assumptions or [])[:30]
        ]
        valid_statuses = {"user_provided", "sourced", "assumption", "estimated"}

        def link_evidence_id(raw_value: Any) -> str:
            value = str(raw_value or "").strip()
            if not value:
                return ""
            value = _evidence_ids([value], 1)[0]
            if value not in linked_evidence:
                if len(linked_evidence) >= 50:
                    raise ValueError("Artifact accepts at most 50 evidence IDs")
                linked_evidence.append(value)
            return value

        def normalize_line(
            row: dict[str, Any],
            *,
            index: int,
            revenue: bool,
        ) -> dict[str, Any]:
            kind = "Revenue" if revenue else "Cost"
            item = _limited_text(str(row.get("item") or ""), 200)
            quantity = _nonnegative_number(row.get("quantity"), f"{kind} {index} quantity")
            price_key = "unit_price" if revenue else "unit_cost"
            unit_price = _nonnegative_number(
                row.get(price_key), f"{kind} {index} {price_key}"
            )
            unit = _limited_text(str(row.get("unit") or ""), 80)
            period = str(row.get("period") or "Unscheduled").strip()[:80]
            status = str(row.get("value_status") or "user_provided").lower()
            if status not in valid_statuses:
                raise ValueError(
                    f"{kind} {index} value_status must be user_provided, "
                    "sourced, assumption, or estimated"
                )
            row_evidence = link_evidence_id(row.get("evidence_id"))
            if status == "sourced" and not row_evidence:
                raise ValueError(f"{kind} {index} requires evidence_id when sourced")
            return {
                "item": item,
                "quantity": quantity,
                "unit": unit,
                "unit_price": unit_price,
                "category": str(
                    row.get("category") or ("revenue" if revenue else "operating")
                )[:80],
                "period": period,
                "value_status": status,
                "evidence_id": row_evidence,
                "total": quantity * unit_price,
            }

        cost_rows = [
            normalize_line(row, index=index, revenue=False)
            for index, row in enumerate(_limited_items(costs, 100), start=1)
        ]
        revenue_rows = [
            normalize_line(row, index=index, revenue=True)
            for index, row in enumerate(_limited_items(revenues, 50), start=1)
        ]
        financing_value = _nonnegative_number(financing_cost, "financing_cost")
        depreciation_value = _nonnegative_number(
            depreciation_cost, "depreciation_cost"
        )
        for item, category, value in (
            ("Financing cost", "financing", financing_value),
            ("Depreciation", "depreciation", depreciation_value),
        ):
            if value:
                cost_rows.append(
                    {
                        "item": item,
                        "quantity": 1.0,
                        "unit": "lump sum",
                        "unit_price": value,
                        "category": category,
                        "period": "Unscheduled",
                        "value_status": "assumption",
                        "evidence_id": "",
                        "total": value,
                    }
                )

        total_cost = sum(row["total"] for row in cost_rows)
        total_revenue = sum(row["total"] for row in revenue_rows)
        net_margin = total_revenue - total_cost
        net_before_financing_tax = (
            total_revenue - total_cost + financing_value + depreciation_value
        )
        margin_percent = (net_margin / total_revenue * 100) if total_revenue else None

        break_even: list[dict[str, Any]] = []
        for row in revenue_rows:
            other_revenue = total_revenue - row["total"]
            required_revenue = max(total_cost - other_revenue, 0.0)
            break_even.append(
                {
                    "item": row["item"],
                    "unit": row["unit"],
                    "break_even_unit_price": (
                        required_revenue / row["quantity"]
                        if row["quantity"] > 0
                        else None
                    ),
                    "break_even_quantity": (
                        required_revenue / row["unit_price"]
                        if row["unit_price"] > 0
                        else None
                    ),
                    "required_revenue_after_other_products": required_revenue,
                }
            )

        cash_flow_by_period: dict[str, dict[str, float]] = {}
        for row in cost_rows:
            if row["category"] == "depreciation":
                continue
            period = row["period"]
            cash_flow_by_period.setdefault(period, {"cash_in": 0.0, "cash_out": 0.0})
            cash_flow_by_period[period]["cash_out"] += row["total"]
        for row in revenue_rows:
            period = row["period"]
            cash_flow_by_period.setdefault(period, {"cash_in": 0.0, "cash_out": 0.0})
            cash_flow_by_period[period]["cash_in"] += row["total"]
        ordered_periods = sorted(
            cash_flow_by_period,
            key=lambda value: (value == "Unscheduled", value.lower()),
        )
        cash_flow: list[dict[str, Any]] = []
        cumulative = 0.0
        for period in ordered_periods:
            values = cash_flow_by_period[period]
            net = values["cash_in"] - values["cash_out"]
            cumulative += net
            cash_flow.append(
                {
                    "period": period,
                    "cash_in": values["cash_in"],
                    "cash_out": values["cash_out"],
                    "net_cash_flow": net,
                    "cumulative_cash_flow": cumulative,
                }
            )

        default_scenarios = [
            {
                "name": "Downside",
                "quantity_change_percent": -10,
                "price_change_percent": -10,
                "cost_change_percent": 10,
            },
            {
                "name": "Base",
                "quantity_change_percent": 0,
                "price_change_percent": 0,
                "cost_change_percent": 0,
            },
            {
                "name": "Upside",
                "quantity_change_percent": 10,
                "price_change_percent": 10,
                "cost_change_percent": -5,
            },
        ]
        scenario_inputs = (sensitivity_scenarios or default_scenarios)[:10]
        sensitivity: list[dict[str, Any]] = []
        for index, scenario in enumerate(scenario_inputs, start=1):
            name = _limited_text(str(scenario.get("name") or f"Scenario {index}"), 80)
            changes: dict[str, float] = {}
            for field in (
                "quantity_change_percent",
                "price_change_percent",
                "cost_change_percent",
            ):
                try:
                    change = float(scenario.get(field) or 0)
                except (TypeError, ValueError) as exc:
                    raise ValueError(f"{name} {field} must be a number") from exc
                if not math.isfinite(change) or change < -100 or change > 1000:
                    raise ValueError(f"{name} {field} must be between -100 and 1000")
                changes[field] = change
            scenario_revenue = (
                total_revenue
                * (1 + changes["quantity_change_percent"] / 100)
                * (1 + changes["price_change_percent"] / 100)
            )
            scenario_cost = total_cost * (1 + changes["cost_change_percent"] / 100)
            sensitivity.append(
                {
                    "name": name,
                    **changes,
                    "total_revenue": scenario_revenue,
                    "total_cost": scenario_cost,
                    "net_margin": scenario_revenue - scenario_cost,
                }
            )
        downside_margin = min(item["net_margin"] for item in sensitivity)

        valid_impact_categories = {
            "water",
            "fertilizer",
            "pesticide",
            "energy",
            "labor",
        }
        normalized_impacts: list[dict[str, Any]] = []
        for index, impact in enumerate((impacts or [])[:50], start=1):
            category = str(impact.get("category") or "").lower()
            if category not in valid_impact_categories:
                raise ValueError(f"Impact {index} has an unsupported category")
            status = str(impact.get("value_status") or "estimated").lower()
            if status not in valid_statuses:
                raise ValueError(f"Impact {index} has an invalid value_status")
            impact_evidence = link_evidence_id(impact.get("evidence_id"))
            if status == "sourced" and not impact_evidence:
                raise ValueError(f"Impact {index} requires evidence_id when sourced")
            normalized_impacts.append(
                {
                    "category": category,
                    "quantity": _nonnegative_number(
                        impact.get("quantity"), f"Impact {index} quantity"
                    ),
                    "unit": _limited_text(str(impact.get("unit") or ""), 80),
                    "period": str(impact.get("period") or "Unscheduled")[:80],
                    "value_status": status,
                    "evidence_id": impact_evidence,
                }
            )

        workbook = Workbook()
        sheet = workbook.active
        sheet.title = "Enterprise budget"
        sheet.append([_safe_cell(_limited_text(title, 200))])
        sheet.append(["Currency", _safe_cell(currency_value)])
        sheet.append(["Geography", _safe_cell(geography_value or "Not supplied")])
        sheet.append(["As of date", effective_date or "Not supplied"])
        sheet.append([])
        sheet.append(["Costs"])
        sheet.append(
            [
                "Item",
                "Quantity",
                "Unit",
                "Unit cost",
                "Category",
                "Period",
                "Value status",
                "Evidence ID",
                "Total cost",
            ]
        )
        cost_start = sheet.max_row + 1
        for row in cost_rows:
            sheet.append(
                [
                    _safe_cell(row["item"]),
                    row["quantity"],
                    _safe_cell(row["unit"]),
                    row["unit_price"],
                    _safe_cell(row["category"]),
                    _safe_cell(row["period"]),
                    row["value_status"],
                    row["evidence_id"],
                    None,
                ]
            )
            current = sheet.max_row
            sheet.cell(current, 9, f"=B{current}*D{current}")
        cost_end = sheet.max_row
        sheet.append(
            ["Total costs", None, None, None, None, None, None, None, f"=SUM(I{cost_start}:I{cost_end})"]
        )
        total_cost_row = sheet.max_row

        sheet.append([])
        sheet.append(["Revenue scenarios"])
        sheet.append(
            [
                "Item",
                "Quantity",
                "Unit",
                "Unit price",
                "Category",
                "Period",
                "Value status",
                "Evidence ID",
                "Total revenue",
            ]
        )
        revenue_start = sheet.max_row + 1
        for row in revenue_rows:
            sheet.append(
                [
                    _safe_cell(row["item"]),
                    row["quantity"],
                    _safe_cell(row["unit"]),
                    row["unit_price"],
                    _safe_cell(row["category"]),
                    _safe_cell(row["period"]),
                    row["value_status"],
                    row["evidence_id"],
                    None,
                ]
            )
            current = sheet.max_row
            sheet.cell(current, 9, f"=B{current}*D{current}")
        revenue_end = sheet.max_row
        sheet.append(
            [
                "Total revenue",
                None,
                None,
                None,
                None,
                None,
                None,
                None,
                f"=SUM(I{revenue_start}:I{revenue_end})",
            ]
        )
        total_revenue_row = sheet.max_row
        sheet.append(
            [
                "Net before financing/depreciation/tax",
                None,
                None,
                None,
                None,
                None,
                None,
                None,
                net_before_financing_tax,
            ]
        )
        sheet.append(
            [
                "Net after financing/depreciation, before tax",
                None,
                None,
                None,
                None,
                None,
                None,
                None,
                f"=I{total_revenue_row}-I{total_cost_row}",
            ]
        )
        for column in "ABCDEFGHI":
            sheet.column_dimensions[column].width = 20
        for row_number in (1, 6, 7, total_cost_row, total_revenue_row):
            for cell in sheet[row_number]:
                cell.font = Font(bold=True)

        break_even_sheet = workbook.create_sheet("Break-even")
        break_even_sheet.append(
            [
                "Revenue item",
                "Unit",
                "Break-even unit price",
                "Break-even quantity",
                "Revenue still required",
            ]
        )
        for row in break_even:
            break_even_sheet.append(
                [
                    _safe_cell(row["item"]),
                    _safe_cell(row["unit"]),
                    row["break_even_unit_price"],
                    row["break_even_quantity"],
                    row["required_revenue_after_other_products"],
                ]
            )

        cash_sheet = workbook.create_sheet("Cash flow")
        cash_sheet.append(
            ["Period", "Cash in", "Cash out", "Net cash flow", "Cumulative cash flow"]
        )
        for row in cash_flow:
            cash_sheet.append(
                [
                    _safe_cell(row["period"]),
                    row["cash_in"],
                    row["cash_out"],
                    row["net_cash_flow"],
                    row["cumulative_cash_flow"],
                ]
            )

        sensitivity_sheet = workbook.create_sheet("Sensitivity")
        sensitivity_sheet.append(
            [
                "Scenario",
                "Quantity change %",
                "Price change %",
                "Cost change %",
                "Revenue",
                "Cost",
                "Net margin",
            ]
        )
        for row in sensitivity:
            sensitivity_sheet.append(
                [
                    _safe_cell(row["name"]),
                    row["quantity_change_percent"],
                    row["price_change_percent"],
                    row["cost_change_percent"],
                    row["total_revenue"],
                    row["total_cost"],
                    row["net_margin"],
                ]
            )

        impact_sheet = workbook.create_sheet("Sustainability impacts")
        impact_sheet.append(
            ["Category", "Quantity", "Unit", "Period", "Value status", "Evidence ID"]
        )
        for row in normalized_impacts:
            impact_sheet.append(
                [
                    row["category"],
                    row["quantity"],
                    _safe_cell(row["unit"]),
                    _safe_cell(row["period"]),
                    row["value_status"],
                    row["evidence_id"],
                ]
            )

        meta = workbook.create_sheet("Assumptions and evidence")
        meta.append(["Generated", datetime.now(UTC).isoformat()])
        meta.append(["Currency", _safe_cell(currency_value)])
        meta.append(["Geography", _safe_cell(geography_value or "Not supplied")])
        meta.append(["As of date", effective_date or "Not supplied"])
        meta.append(["Assumptions"])
        for item in assumption_rows:
            meta.append([_safe_cell(item)])
        meta.append(["Evidence IDs"])
        for item in linked_evidence:
            meta.append([item])
        meta.append(["Legacy sources"])
        for item in legacy_sources[:30]:
            meta.append([_safe_cell(item)])

        buffer = io.BytesIO()
        workbook.save(buffer)
        filename = _safe_filename(title, ".xlsx")
        statuses = [row["value_status"] for row in [*cost_rows, *revenue_rows]]
        status_counts = {status: statuses.count(status) for status in sorted(set(statuses))}
        provenance = _provenance_metadata(linked_evidence, legacy_sources)
        warnings = ["Scenario estimate; not a profit guarantee."]
        if not geography_value:
            warnings.append("Geography was not supplied.")
        if not effective_date:
            warnings.append("As-of date was not supplied.")
        if provenance["provenance_status"] != "evidence_linked":
            warnings.append("Material values are not linked to immutable evidence IDs.")
        if financing_value or depreciation_value:
            warnings.append("Financing and depreciation values are scenario assumptions.")
        reference = self._save(
            "enterprise_budget",
            filename,
            XLSX_MIME,
            buffer.getvalue(),
            {
                "currency": currency_value,
                "geography": geography_value,
                "as_of_date": effective_date,
                "total_cost": total_cost,
                "total_revenue": total_revenue,
                "net_before_financing_tax": net_before_financing_tax,
                "net_after_financing_depreciation_before_tax": net_margin,
                "downside_margin": downside_margin,
                "input_status_counts": status_counts,
                "impact_count": len(normalized_impacts),
                **provenance,
            },
        )
        return {
            **reference.to_dict(),
            "currency": currency_value,
            "geography": geography_value,
            "as_of_date": effective_date,
            "total_cost": round(total_cost, 2),
            "total_revenue": round(total_revenue, 2),
            "net_before_financing_tax": round(net_before_financing_tax, 2),
            "net_after_financing_depreciation_before_tax": round(net_margin, 2),
            "expected_margin": round(net_margin, 2),
            "margin_percent": round(margin_percent, 2) if margin_percent is not None else None,
            "break_even": break_even,
            "cash_flow": cash_flow,
            "sensitivity": sensitivity,
            "downside_margin": round(downside_margin, 2),
            "impacts": normalized_impacts,
            "input_status_counts": status_counts,
            **provenance,
            "warnings": warnings,
            "warning": warnings[0],
        }


def convert_agricultural_units(
    value: float,
    from_unit: str,
    to_unit: str,
) -> dict[str, Any]:
    units = {
        "m2": ("area", 1.0),
        "dunam": ("area", 1000.0),
        "hectare": ("area", 10000.0),
        "liter": ("volume", 1.0),
        "m3": ("volume", 1000.0),
        "kg": ("mass", 1.0),
        "tonne": ("mass", 1000.0),
    }
    source = from_unit.lower().strip()
    target = to_unit.lower().strip()
    if source not in units or target not in units:
        raise ValueError("Unsupported unit")
    source_dimension, source_factor = units[source]
    target_dimension, target_factor = units[target]
    if source_dimension != target_dimension:
        raise ValueError("Units measure different dimensions")
    result = float(value) * source_factor / target_factor
    return {
        "value": float(value),
        "from_unit": source,
        "result": round(result, 8),
        "to_unit": target,
    }

"""Convert the supplied DOCX into the canonical bilingual pilot Markdown corpus."""

from __future__ import annotations

import argparse
import hashlib
import json
import re
import sys
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any
from urllib.parse import urlsplit, urlunsplit

from docx import Document
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph

from farmers_chatbot.agrifood_ontology import (
    ONTOLOGY_VERSION,
    ontology_for_record,
    validate_ontology,
)

try:
    from scripts.agrifood_arabic_drafts import (
        ARABIC_DRAFTS,
        validate_local_drafts,
    )
except ModuleNotFoundError:  # direct script entry point
    from agrifood_arabic_drafts import (  # type: ignore[no-redef]
        ARABIC_DRAFTS,
        validate_local_drafts,
    )

SOURCE_SHA256 = "3C0BAF8145E4A2E287BA5783F8AB26A7CEAE1C5340322C1EE065887CC9B75B0E"
GENERATED_AT = "2026-08-11T00:00:00+00:00"
SOURCE_RE = re.compile(r"\[(S\d{2})\]\s*(.*)")
CHAPTER_RE = re.compile(r"Chapter\s+(\d+)\s+[—-]\s+(.+)", re.IGNORECASE)
SECTION_RE = re.compile(r"^([A-M])\.\s*(.+)")
URL_RE = re.compile(r"https?://\S+")
ARABIC_RE = re.compile(r"[\u0600-\u06ff]")
OPAQUE_SOURCE_GROUP_RE = re.compile(r"\[([^\]]*S\d{2}[^\]]*)\]")
RETAINED_SECTIONS = {"C", "D", "E", "G", "H"}

DOC_SOURCE_IDS = {
    "S01": "ESDU-HOME-2026", "S02": "ESDU-ABOUT-2026",
    "S03": "ESDU-ISNAD", "S04": "ESDU-AKKAR-VALUECHAINS",
    "S05": "ESDU-CLIMAT-AKKAR", "S06": "ESDU-ARDI-ARDAK",
    "S07": "ESDU-KARIANET", "S08": "ESDU-RESOLVE",
    "S09": "MOA-LEBANON-NAS-2020-2025",
    "S10": "FAO-LEBANON-RESILIENT-LIVELIHOODS",
    "S11": "FAO-CROP-EVAPOTRANSPIRATION-56-2025",
    "S12": "FAO-WATER-QUALITY-1985", "S13": "FAO-SOIL-TESTING-2019",
    "S14": "FAO-WHO-PESTICIDE-CODE-2014",
    "S15": "CODEX-FOOD-HYGIENE-CXC1-2022",
    "S16": "WHO-FIVE-KEYS-SAFER-FOOD-2006",
    "S17": "WHO-GROWING-SAFER-PRODUCE-2012",
    "S18": "WOAH-CODES-MANUALS-2026",
    "S19": "WOAH-ANTIMICROBIAL-USE-2024",
    "S20": "FAO-CLIMATE-LIVESTOCK-2023",
    "S21": "FAO-POSTHARVEST-GRAIN-1996",
    "S22": "WMO-CLIMATE-SERVICES-2026",
    "S23": "FAO-GREENHOUSE-GAP-2013", "S24": "FAO-SAVE-GROW-2011",
}

LEGACY_JSON_EXCLUSIONS = {
    "ESDU-IDENTITY-010": (
        "institutional identity belongs in product provenance, not farmer guidance"
    ),
    "ESDU-LIVING-LABS-011": (
        "institutional outcome claims require passage and institutional approval"
    ),
    "ESDU-LOCAL-FOOD-012": (
        "institutional outcome/service claims require passage and institutional approval"
    ),
    "LEGUMES-RELEAF-013": (
        "project-specific outcomes and active opportunities require current "
        "official validation"
    ),
}


@dataclass(frozen=True)
class RecordSpec:
    record_id: str
    title_en: str
    title_ar: str
    chapters: tuple[int, ...]
    guide_ids: tuple[str, ...]
    topics: tuple[str, ...]
    risk: str = "medium"
    kind: str = "evidence"
    dynamicity: str = "stable"
    extra_sources: tuple[str, ...] = ()
    relations: tuple[tuple[str, str], ...] = ()


SPECS = (
    RecordSpec("kb-scope-local-context", "Akkar context and limits on transferring advice", "سياق عكار وحدود نقل الإرشادات", (3, 4, 8), ("AKKAR-PROFILE-001", "AKKAR-SECTOR-002", "FARMER-QUESTION-018"), ("Akkar", "local context", "applicability"), extra_sources=("MOA-AKKAR-2026", "UNDP-LEBANON-NAP-2025"), relations=(("requires_context", "kb-decision-rules"),)),
    RecordSpec("kb-crop-production", "Crop production decisions", "قرارات إنتاج المحاصيل", (11,), ("POTATO-DECISIONS-003", "ORCHARD-DECISIONS-005"), ("crop", "potato", "orchard", "production stage"), relations=(("depends_on", "kb-soil-management"), ("depends_on", "kb-water-irrigation"))),
    RecordSpec("kb-livestock", "Livestock and mixed-farm decisions", "قرارات الثروة الحيوانية والمزارع المختلطة", (12,), ("LIVESTOCK-SYSTEMS-008",), ("animal", "livestock", "mixed farm"), risk="high", relations=(("escalates_to", "kb-referrals"),)),
    RecordSpec("kb-soil-management", "Soil observation, testing, and fertility decisions", "قرارات معاينة التربة وفحصها وخصوبتها", (13,), ("SOIL-FERTILITY-007",), ("soil", "fertility", "soil testing"), relations=(("supports_action", "kb-crop-production"),)),
    RecordSpec("kb-water-irrigation", "Water and irrigation decisions", "قرارات المياه والري", (14,), ("WATER-IRRIGATION-006",), ("water", "irrigation", "water quality"), extra_sources=("UNDP-IRRIGATION-AKKAR",), relations=(("depends_on", "kb-climate-season"),)),
    RecordSpec("kb-ipm-safety", "Integrated pest management and pesticide safety", "الإدارة المتكاملة للآفات وسلامة المبيدات", (15, 16), ("PESTICIDE-VERIFICATION-019", "WORKER-CHILD-SAFETY-020", "SAFE-ADVICE-016"), ("pest", "IPM", "pesticide", "worker safety"), risk="high", kind="policy", dynamicity="live_only", extra_sources=("MOA-REGISTERED-PESTICIDES-2026", "MOA-BANNED-PESTICIDES-2026", "FAO-LEBANON-PESTICIDE-CHILD-SAFETY"), relations=(("requires_live_source", "kb-dynamic-information"), ("escalates_to", "kb-referrals"))),
    RecordSpec("kb-climate-season", "Climate and seasonal decision support", "دعم القرارات المناخية والموسمية", (17,), ("CLIMATE-RISK-014",), ("climate", "season", "calendar", "adaptation"), dynamicity="live_only", extra_sources=("UNDP-LEBANON-NAP-2025",), relations=(("requires_live_source", "kb-dynamic-information"),)),
    RecordSpec("kb-greenhouse", "Greenhouse crop decisions", "قرارات محاصيل البيوت المحمية", (18,), ("GREENHOUSE-DECISIONS-004",), ("practice", "greenhouse", "tomato", "cucumber"), extra_sources=("MOA-AKKAR-GREENHOUSE-2026",), relations=(("depends_on", "kb-water-irrigation"), ("depends_on", "kb-ipm-safety"))),
    RecordSpec("kb-postharvest", "Harvest and post-harvest decisions", "قرارات الحصاد وما بعد الحصاد", (19,), ("POSTHARVEST-MARKET-015",), ("production_stage", "harvest", "storage", "quality"), relations=(("supports_action", "kb-business-markets"),)),
    RecordSpec("kb-food-processing-safety", "Food processing and food-safety boundaries", "حدود تصنيع الغذاء وسلامته", (20,), (), ("risk", "processing", "food safety"), risk="high", relations=(("escalates_to", "kb-referrals"),)),
    RecordSpec("kb-business-markets", "Farm business, value-chain, and market decisions", "قرارات الأعمال الزراعية وسلاسل القيمة والأسواق", (21,), ("VALUE-CHAINS-009", "POSTHARVEST-MARKET-015"), ("market", "business", "cost", "value chain"), dynamicity="live_only", relations=(("requires_live_source", "kb-dynamic-information"),)),
    RecordSpec("kb-troubleshooting", "Troubleshooting without premature diagnosis", "استكشاف المشكلات دون تشخيص متسرّع", (22,), ("FARMER-QUESTION-018", "SAFE-ADVICE-016"), ("symptom", "diagnosis", "observation"), risk="high", relations=(("may_be_confused_with", "kb-ipm-safety"), ("escalates_to", "kb-referrals"))),
    RecordSpec("kb-decision-rules", "Recommendation and decision rules", "قواعد التوصية واتخاذ القرار", (23,), ("FARMER-QUESTION-018", "SAFE-ADVICE-016"), ("practice", "decision rules", "context", "uncertainty"), kind="workflow", relations=(("requires_context", "kb-scope-local-context"),)),
    RecordSpec("kb-faq", "Frequently asked farm questions", "الأسئلة الزراعية الشائعة", (25,), (), ("practice", "FAQ", "farmer questions"), kind="workflow", relations=(("related_to", "kb-decision-rules"),)),
    RecordSpec("kb-misconceptions", "Common agrifood misconceptions", "مفاهيم زراعية وغذائية شائعة وخاطئة", (26,), (), ("risk", "misconceptions", "evidence"), relations=(("conflicts_with", "kb-decision-rules"),)),
    RecordSpec("kb-referrals", "Useful referral and escalation hand-offs", "الإحالات والتصعيد العملي المفيد", (27,), ("EXTENSION-REFERRAL-021", "SAFE-ADVICE-016"), ("service", "referral", "extension", "expert"), kind="workflow", dynamicity="live_only", extra_sources=("MOA-EXTENSION-CENTERS",), relations=(("requires_live_source", "kb-dynamic-information"),)),
    RecordSpec("kb-terminology", "Arabic, English, and local agrifood terminology", "المصطلحات الزراعية والغذائية العربية والإنجليزية والمحلية", (28,), (), ("practice", "terminology", "Arabic", "Lebanese"), risk="low", kind="glossary", relations=(("related_to", "kb-scope-local-context"),)),
    RecordSpec("kb-dynamic-information", "Information that requires a dated live source", "المعلومات التي تتطلب مصدراً حياً مؤرخاً", (7,), ("DYNAMIC-INFORMATION-017",), ("risk", "weather", "prices", "regulation"), kind="policy", dynamicity="live_only", relations=(("requires_live_source", "kb-referrals"),)),
)


@dataclass
class DraftRecord:
    spec: RecordSpec
    english_guidance: str
    decision_logic: str
    safe_next_action: str
    avoid_escalate: str
    applicability_limits: str
    source_ids: tuple[str, ...]
    supersedes: tuple[str, ...]
    translation: dict[str, str] = field(default_factory=dict)


def normalize(value: str) -> str:
    return re.sub(r"\s+", " ", value).strip()


def url_key(value: str) -> str:
    parts = urlsplit(value.rstrip(".,;)"))
    return urlunsplit((parts.scheme.casefold(), parts.netloc.casefold().removeprefix("www."), parts.path.rstrip("/"), "", ""))


def source_ids_in_text(value: str) -> set[str]:
    """Expand single, list, and numeric-range DOCX source references."""

    opaque_ids: set[str] = set()
    for match in OPAQUE_SOURCE_GROUP_RE.finditer(value):
        group = match.group(1)
        tokens = re.findall(r"S\d{2}", group)
        if len(tokens) == 2 and re.search(r"[–—-]", group):
            start, end = (int(token[1:]) for token in tokens)
            if start <= end:
                tokens = [f"S{number:02d}" for number in range(start, end + 1)]
        opaque_ids.update(token for token in tokens if token in DOC_SOURCE_IDS)
    return {DOC_SOURCE_IDS[item] for item in opaque_ids}


def replace_source_ids(value: str) -> str:
    def replace(match: re.Match[str]) -> str:
        group = match.group(0)
        source_ids = sorted(source_ids_in_text(group))
        if not source_ids:
            return group
        label = "source" if len(source_ids) == 1 else "sources"
        return f"[{label}: {', '.join(source_ids)}]"

    value = OPAQUE_SOURCE_GROUP_RE.sub(replace, value)
    value = value.replace("[VERIFIED EXTERNAL SOURCE]", "")
    return value.replace("[DATE-SENSITIVE]", "(date-sensitive)")


def table_markdown(table: Table) -> str:
    rows: list[list[str]] = []
    for row in table.rows:
        values: list[str] = []
        for cell in row.cells:
            text = normalize(cell.text)
            if not values or values[-1] != text:
                values.append(text)
        if values and any(values):
            rows.append(values)
    if not rows:
        return ""
    width = max(map(len, rows))
    rows = [row + [""] * (width - len(row)) for row in rows]
    if width == 1:
        return re.sub(r"^(SAFETY BOUNDARY|DRAFT STATUS)\s*", "", rows[0][0], flags=re.IGNORECASE)
    lines = ["| " + " | ".join(rows[0]) + " |", "| " + " | ".join(["---"] * width) + " |"]
    return "\n".join([*lines, *("| " + " | ".join(row) + " |" for row in rows[1:])])


def allowed_text(value: str) -> bool:
    lowered = value.casefold()
    blocked = ("esdu knowledge opportunity", "knowledge gap]", "draft completeness", "ai-assisted working draft", "rag test question", "expert review required", "suggested data collection")
    return not any(item in lowered for item in blocked) and "esdu" not in lowered and "american university of beirut" not in lowered


def extract_chapters(path: Path) -> tuple[dict[int, dict[str, list[str]]], dict[int, set[str]]]:
    document = Document(path)
    chapters: dict[int, dict[str, list[str]]] = {}
    source_ids: dict[int, set[str]] = {}
    chapter: int | None = None
    section: str | None = None
    for child in document.element.body.iterchildren():
        if child.tag == qn("w:p"):
            paragraph = Paragraph(child, document)
            text = normalize(paragraph.text)
            if not text:
                continue
            match = CHAPTER_RE.fullmatch(text)
            if match and paragraph.style and paragraph.style.name.startswith("Heading 1"):
                chapter, section = int(match.group(1)), None
                chapters.setdefault(chapter, {})
                source_ids.setdefault(chapter, set())
                continue
            if chapter is None:
                continue
            section_match = SECTION_RE.match(text)
            if paragraph.style and paragraph.style.name.startswith("Heading 2"):
                section = section_match.group(1) if section_match else None
                continue
            if section in RETAINED_SECTIONS and allowed_text(text):
                source_ids[chapter].update(source_ids_in_text(text))
                chapters[chapter].setdefault(section, []).append(replace_source_ids(text))
        elif child.tag == qn("w:tbl") and chapter is not None and section in RETAINED_SECTIONS:
            value = table_markdown(Table(child, document))
            if value and allowed_text(value):
                source_ids[chapter].update(source_ids_in_text(value))
                chapters[chapter].setdefault(section, []).append(replace_source_ids(value))
    return chapters, source_ids


def dedupe_blocks(blocks: list[str], seen: set[str]) -> str:
    output = []
    for block in blocks:
        key = re.sub(r"[^\w\u0600-\u06ff]+", " ", block.casefold()).strip()
        if key and key not in seen:
            seen.add(key)
            output.append(block)
    return "\n\n".join(output)


def build_records(chapters: dict[int, dict[str, list[str]]], chapter_sources: dict[int, set[str]], guide_path: Path) -> list[DraftRecord]:
    guide = json.loads(guide_path.read_text(encoding="utf-8"))
    guide_by_id = {item["id"]: item for item in guide["items"]}
    records, seen = [], set()
    owned_legacy_items: set[str] = set()
    defaults = {
        "C": "No additional standalone guidance was retained; use the linked decision and safety sections.",
        "D": "Collect the missing local context and use the lowest-risk reversible next step.",
        "E": "If evidence or context is insufficient, state the gap and do not infer a precise recommendation.",
        "G": "Apply only within the stated geography, production system, season, and evidence limits.",
        "H": "Escalate when the decision could affect human, animal, food, environmental, or legal safety.",
    }
    for spec in SPECS:
        blocks = {key: [] for key in RETAINED_SECTIONS}
        sources, supersedes = set(spec.extra_sources), []
        for guide_id in spec.guide_ids:
            item = guide_by_id[guide_id]
            if guide_id not in owned_legacy_items:
                supersedes.append(guide_id)
                owned_legacy_items.add(guide_id)
            sources.update(item.get("source_ids") or [])
            if allowed_text(item["text_en"]):
                blocks["C"].append(normalize(item["text_en"]))
        for chapter in spec.chapters:
            sources.update(chapter_sources.get(chapter, set()))
            for section in RETAINED_SECTIONS:
                blocks[section].extend(chapters.get(chapter, {}).get(section, []))
        values = {key: dedupe_blocks(blocks[key], seen) or defaults[key] for key in RETAINED_SECTIONS}
        if spec.record_id == "kb-decision-rules":
            values["G"] = (
                "Apply these decision rules only within the declared geography, "
                "production system, season, evidence scope, and channel policy."
            )
        elif spec.record_id == "kb-misconceptions":
            values["E"] = (
                "When evidence or context is insufficient to correct a claim, "
                "state the gap and do not infer a precise recommendation."
            )
        elif spec.record_id == "kb-terminology":
            values["G"] = (
                "Use each term only for its reviewed geography, language "
                "community, production system, and technical context."
            )
        records.append(DraftRecord(spec, values["C"], values["E"], values["D"], values["H"], values["G"], tuple(sorted(sources)), tuple(sorted(set(supersedes)))))
    return records


def apply_local_arabic_drafts(
    records: list[DraftRecord],
    guide_path: Path,
) -> None:
    """Attach repository-owned Arabic text without calling a network provider."""

    guide = json.loads(guide_path.read_text(encoding="utf-8"))
    guide_by_id = {item["id"]: item for item in guide["items"]}
    validate_local_drafts({record.spec.record_id for record in records})
    for record in records:
        existing: list[str] = []
        seen: set[str] = set()
        for guide_id in record.spec.guide_ids:
            text = normalize(str(guide_by_id[guide_id].get("text_ar") or ""))
            key = re.sub(r"\W+", " ", text.casefold()).strip()
            if text and allowed_text(text) and key not in seen:
                existing.append(text)
                seen.add(key)
        local = {
            key: value.strip()
            for key, value in ARABIC_DRAFTS[record.spec.record_id].items()
        }
        record.translation = {
            "title_ar": record.spec.title_ar,
            **local,
            "guidance_ar": "\n\n".join([*existing, local["guidance_ar"]]),
        }


def doc_sources(path: Path) -> dict[str, dict[str, Any]]:
    result = {}
    for paragraph in Document(path).paragraphs:
        match = SOURCE_RE.fullmatch(normalize(paragraph.text))
        if not match or match.group(1) in result:
            continue
        opaque, description = match.groups()
        url_match = URL_RE.search(description)
        url = url_match.group(0).rstrip(".,") if url_match else None
        title = description[: url_match.start()].strip() if url_match else description
        result[opaque] = {"id": DOC_SOURCE_IDS[opaque], "legacy_ids": [opaque], "title": title, "publisher": title.split(".", 1)[0], "url": url, "source_class": "A", "review_status": "official_public_source", "production_eligible": False}
    return result


def reconcile_sources(docx_path: Path, sources_path: Path) -> dict[str, dict[str, Any]]:
    current = json.loads(sources_path.read_text(encoding="utf-8"))
    merged = {item["id"]: dict(item) for item in current}
    urls = {url_key(item["url"]): item["id"] for item in current if item.get("url")}
    for opaque, item in doc_sources(docx_path).items():
        target_id = urls.get(url_key(item["url"])) if item.get("url") else None
        target_id = target_id or DOC_SOURCE_IDS[opaque]
        if target_id in merged:
            merged[target_id]["legacy_ids"] = sorted(set(merged[target_id].get("legacy_ids") or []) | {opaque})
        else:
            item["id"] = target_id
            merged[target_id] = item
    return merged


def metadata(record: DraftRecord) -> dict[str, Any]:
    spec = record.spec
    ontology = ontology_for_record(spec.record_id)
    return {
        "id": spec.record_id, "title_en": spec.title_en,
        "title_ar": record.translation["title_ar"], "languages": ["en", "ar"],
        "content_kind": spec.kind, "geography": ["Akkar", "rural Lebanon"],
        "topics": list(spec.topics),
        "entities": ontology["ontology_entities"],
        "graph_relations": [{"type": relation, "target": target} for relation, target in spec.relations],
        **ontology,
        "risk": spec.risk, "evidence_class": "official_and_draft_synthesis",
        "claim_ids": [f"claim:{spec.record_id}:guidance", f"claim:{spec.record_id}:decision", f"claim:{spec.record_id}:safety"],
        "source_ids": list(record.source_ids), "review_status": "ai_draft",
        "translation_status": "machine_draft", "dynamicity": spec.dynamicity,
        "translation_method": "local_repository_ai_draft",
        "effective_from": None, "expires_at": None, "review_by": "2026-11-11",
        "owner_role": "knowledge_steward",
        "reviewer_roles": ["domain_expert", "Arabic_reviewer", "field_reviewer"],
        "retrieval_enabled": True, "publication_scope": "pilot",
        "production_eligible": False, "supersedes_legacy_items": list(record.supersedes),
    }


def render_markdown(records: list[DraftRecord], sources: dict[str, dict[str, Any]], source_name: str) -> str:
    lines = [
        "---", "document_id: raise-agrifood-knowledge", 'version: "0.2"',
        "status: ai_draft", "publication_scope: pilot", "production_eligible: false",
        f"source_doc_sha256: {SOURCE_SHA256}", f"generated_at: {GENERATED_AT}",
        "languages: [en, ar]", "scope: [Akkar, rural Lebanon]",
        "default_review_status: ai_draft", "default_translation_status: machine_draft",
        "translation_method: local_repository_ai_draft",
        "default_retrieval_enabled: true", f"ontology_version: {ONTOLOGY_VERSION}",
        "---", "", "# RAISE Agrifood Knowledge Draft", "",
        "> Pilot draft: English is the authoritative draft source. Arabic was drafted locally in the repository without an external translation service. Neither language is production-approved; high-risk details remain limited or referral-only until expert review.", "",
    ]
    for record in records:
        lines.extend([
            f"## {record.spec.title_en} / {record.translation['title_ar']}", "", "```yaml",
            json.dumps(metadata(record), ensure_ascii=False, indent=2, sort_keys=True), "```", "",
            "### English guidance", "", record.english_guidance, "",
            "### Arabic guidance — machine draft", "", record.translation["guidance_ar"], "",
            "### Decision logic", "", record.decision_logic, "",
            "### منطق القرار — مسودة آلية", "", record.translation["decision_logic_ar"], "",
            "### Safe next action", "", record.safe_next_action, "",
            "### الخطوة التالية الآمنة — مسودة آلية", "", record.translation["safe_next_action_ar"], "",
            "### Avoid or escalate", "", record.avoid_escalate, "",
            "### ما يجب تجنبه أو تصعيده — مسودة آلية", "", record.translation["avoid_escalate_ar"], "",
            "### Evidence and applicability limits", "", record.applicability_limits, "",
            "### حدود الأدلة وقابلية التطبيق — مسودة آلية", "",
            record.translation["applicability_limits_ar"], "", "### Claim-level sources", "",
        ])
        for source_id in record.source_ids:
            source = sources[source_id]
            lines.append(f"- [{source_id}] {source.get('title', source_id)} — {source.get('url') or 'URL unresolved; validation required'}")
        lines.append("")

    retained = sorted({chapter for spec in SPECS for chapter in spec.chapters})
    legacy_owners = {
        legacy_id: record.spec.record_id
        for record in records
        for legacy_id in record.supersedes
    }
    mapping = {
        "retained_and_merged": {str(chapter): [spec.record_id for spec in SPECS if chapter in spec.chapters] for chapter in retained},
        "excluded_from_retrieval": {
            "1-2": "product purpose and interface material belongs in product documentation",
            "5": "unmeasured product differentiation claim",
            "6": "source authority rules consolidated in this appendix and ingestion validation",
            "9-10": "institutional service and outcome claims pending institutional approval",
            "24": "institution-specific case claims pending passage and institutional review",
            "29-30": "provenance rules and gaps consolidated in this appendix",
            "31": "benchmark material moved to evaluation fixtures",
            "32": "governance and change control belongs in operational documentation",
            "appendices": "review registers, prompts, dashboards, and readiness checklists are non-retrievable",
        },
        "legacy_json_merge_owners": dict(sorted(legacy_owners.items())),
        "legacy_json_excluded": LEGACY_JSON_EXCLUSIONS,
    }
    source_register = []
    for source_id in sorted(sources):
        item = dict(sources[source_id])
        item.setdefault("review_status", "official_public_source")
        item.update({"retrieval_enabled": False, "production_eligible": False})
        source_register.append(item)
    lines.extend([
        "# Non-retrievable source and validation appendix", "",
        "This appendix is excluded from retrieval. It records provenance, mapping, and unresolved approval work.", "",
        "### Evidence and source policy", "",
        "- Prefer current primary official or peer-reviewed passages appropriate to the claim and geography.",
        "- A domain, publisher, or evidence class does not by itself verify a claim; the retained passage must directly support it.",
        "- Separate evidence strength from local applicability, retrieval confidence, review status, and production eligibility.",
        "- Record conflicts, supersession, observation time, expiry, geography, and material limitations instead of silently choosing a source.",
        "- Treat missing passage support as an explicit validation gap; medium/high/critical actions remain limited, refused, or referred as policy requires.",
        "- Dynamic facts remain live evidence and are never promoted into the permanent graph as undated values.", "",
        "### Source register", "", "```json",
        json.dumps(source_register, ensure_ascii=False, indent=2, sort_keys=True), "```", "",
        "### Conversion and merge map", "", "```json",
        json.dumps(mapping, ensure_ascii=False, indent=2, sort_keys=True), "```", "",
        "### Validation backlog", "",
        "- Obtain expert and field review for every retained draft record and its Arabic wording.",
        "- Verify passage-level entailment, license, archive hash, geography, dates, and supersession for each source.",
        "- Keep current prices, weather, alerts, grants, contacts, pesticide registers, and regulations live-only with observation and expiry times.",
        "- Do not add chemical doses, veterinary treatment, definitive diagnosis, food-processing parameters, or local agronomic thresholds without approved evidence.",
        "- Treat the product restriction on glyphosate recommendations as RAISE policy, not a legal-status claim.",
        "- Confirm institutional services, project outputs, outcomes, contacts, and referral availability before retrieval use.",
        "- Run hidden bilingual retrieval, citation, graph, safety, and farmer/expert preference evaluations before release activation.", "",
        f"Source converted from `{source_name}`; the source remains unchanged at SHA-256 `{SOURCE_SHA256}`.", "",
    ])
    return "\n".join(lines)


def render_arabic_companion(
    records: list[DraftRecord],
    sources: dict[str, dict[str, Any]],
    source_name: str,
    canonical_name: str,
) -> str:
    """Render a standalone Arabic review view with identical graph IDs."""

    lines = [
        "---",
        "document_id: raise-agrifood-knowledge-ar",
        'version: "0.2"',
        "status: ai_draft",
        "publication_scope: pilot",
        "production_eligible: false",
        f"source_doc_sha256: {SOURCE_SHA256}",
        f"generated_at: {GENERATED_AT}",
        "languages: [ar]",
        "scope: [Akkar, rural Lebanon]",
        f"canonical_companion: {canonical_name}",
        "translation_status: machine_draft",
        "translation_method: local_repository_ai_draft",
        f"ontology_version: {ONTOLOGY_VERSION}",
        "---",
        "",
        "# مسودة RAISE للمعرفة الزراعية والغذائية",
        "",
        (
            "> هذه نسخة عربية تجريبية أُعدّت محلياً داخل المستودع من دون "
            "إرسال نص الوثيقة إلى خدمة ترجمة خارجية. تبقى الإنجليزية في "
            "الملف الأساسي هي المسودة المرجعية إلى حين المراجعة الزراعية "
            "واللغوية والميدانية."
        ),
        "",
    ]
    for record in records:
        lines.extend(
            [
                f"## {record.translation['title_ar']}",
                "",
                "~~~yaml",
                json.dumps(
                    {
                        **metadata(record),
                        "view_language": "ar",
                        "canonical_record_id": record.spec.record_id,
                    },
                    ensure_ascii=False,
                    indent=2,
                    sort_keys=True,
                ),
                "~~~",
                "",
                "### الإرشادات — مسودة آلية محلية",
                "",
                record.translation["guidance_ar"],
                "",
                "### منطق القرار — مسودة آلية محلية",
                "",
                record.translation["decision_logic_ar"],
                "",
                "### الخطوة التالية الآمنة — مسودة آلية محلية",
                "",
                record.translation["safe_next_action_ar"],
                "",
                "### ما يجب تجنبه أو تصعيده — مسودة آلية محلية",
                "",
                record.translation["avoid_escalate_ar"],
                "",
                "### حدود الأدلة وقابلية التطبيق — مسودة آلية محلية",
                "",
                record.translation["applicability_limits_ar"],
                "",
                "### المصادر على مستوى الادعاء",
                "",
            ]
        )
        for source_id in record.source_ids:
            source = sources[source_id]
            lines.append(
                f"- [{source_id}] {source.get('title', source_id)} — "
                f"{source.get('url') or 'الرابط غير محسوم؛ يلزم التحقق'}"
            )
        lines.append("")

    source_register = []
    for source_id in sorted(sources):
        item = dict(sources[source_id])
        item.setdefault("review_status", "official_public_source")
        item.update({"retrieval_enabled": False, "production_eligible": False})
        source_register.append(item)
    lines.extend(
        [
            "# ملحق المصادر والتحقق غير القابل للاسترجاع",
            "",
            "هذا الملحق خارج الاسترجاع ويحتفظ بسجل المصادر وحالة المراجعة.",
            "",
            "### سياسة الأدلة والمصادر",
            "",
            "- تُفضّل المقاطع الأولية الرسمية الحديثة أو المحكمة والملائمة للادعاء والجغرافيا.",
            "- لا يثبت النطاق أو الناشر أو فئة الدليل الادعاء وحده؛ يجب أن يدعمه المقطع المحفوظ مباشرة.",
            "- تُفصل قوة الدليل عن الملاءمة المحلية وثقة الاسترجاع وحالة المراجعة وأهلية الإنتاج.",
            "- يُسجّل التعارض والنسخة الأحدث ووقت المشاهدة والانتهاء والجغرافيا والحدود بدلاً من اختيار مصدر بصمت.",
            "- يُعامل غياب دعم المقطع كفجوة تحقق صريحة، وتبقى الأفعال المتوسطة أو العالية أو الحرجة محدودة أو مرفوضة أو محالة بحسب السياسة.",
            "- تبقى الحقائق المتغيرة أدلة حية ولا تُضاف إلى الرسم الدائم كقيم غير مؤرخة.",
            "",
            "### سجل المصادر",
            "",
            "~~~json",
            json.dumps(
                source_register,
                ensure_ascii=False,
                indent=2,
                sort_keys=True,
            ),
            "~~~",
            "",
            "### أعمال التحقق المتبقية",
            "",
            (
                "- مراجعة كل سجل وصياغته العربية من خبير مجال ومراجع لغوي "
                "ومستخدمين ممثلين."
            ),
            (
                "- التحقق من إسناد كل ادعاء إلى مقطع مصدر ومن الترخيص "
                "والتاريخ والجغرافيا والنسخة الأحدث."
            ),
            (
                "- إبقاء الأسعار والطقس والتنبيهات والمنح والاتصالات وسجلات "
                "المبيدات والقوانين ضمن مصادر حية مؤرخة."
            ),
            (
                "- عدم إضافة جرعات كيميائية أو علاج بيطري أو تشخيص نهائي أو "
                "معاملات تصنيع غذائي دقيقة من دون دليل معتمد."
            ),
            "",
            (
                f"حُوّل المصدر من '{source_name}' وبقي دون تعديل عند "
                f"SHA-256 '{SOURCE_SHA256}'."
            ),
            "",
        ]
    )
    return "\n".join(lines)


def validate_arabic_companion(
    text: str,
    records: list[DraftRecord],
    sources: dict[str, Any],
) -> None:
    if "\ufffd" in text or text.count("\n## ") != len(records):
        raise ValueError("Arabic companion structure or Unicode is invalid")
    if "# ملحق المصادر والتحقق غير القابل للاسترجاع" not in text:
        raise ValueError("Arabic non-retrievable appendix is missing")
    record_ids = {record.spec.record_id for record in records}
    metadata_blocks = [
        json.loads(value)
        for value in re.findall(
            r"(?ms)^## .+?\r?$.*?^~~~yaml\r?\n(.*?)\r?\n~~~",
            text,
        )
    ]
    parsed_ids = [str(item.get("id") or "") for item in metadata_blocks]
    if len(metadata_blocks) != len(records) or set(parsed_ids) != record_ids:
        raise ValueError("Arabic companion metadata IDs are incomplete")
    if len(parsed_ids) != len(set(parsed_ids)):
        raise ValueError("Arabic companion metadata IDs are duplicated")
    for item in metadata_blocks:
        if item.get("view_language") != "ar":
            raise ValueError("Arabic companion has a non-Arabic record view")
        if item.get("production_eligible") is not False:
            raise ValueError("Arabic companion record is production eligible")
        if item.get("review_status") != "ai_draft":
            raise ValueError("Arabic companion record is not an AI draft")
        if set(item.get("source_ids") or []) - set(sources):
            raise ValueError(f"Arabic sources are unresolved for {item['id']}")
        for relation in item.get("graph_relations") or []:
            if relation.get("target") not in record_ids:
                raise ValueError(
                    f"Arabic graph target is unresolved: {relation.get('target')}"
                )
    for record in records:
        if f'"id": "{record.spec.record_id}"' not in text:
            raise ValueError(f"Arabic companion is missing {record.spec.record_id}")
        if any(
            not ARABIC_RE.search(record.translation[field])
            for field in (
                "guidance_ar",
                "decision_logic_ar",
                "safe_next_action_ar",
                "avoid_escalate_ar",
                "applicability_limits_ar",
            )
        ):
            raise ValueError(
                f"Arabic text is incomplete for {record.spec.record_id}"
            )
        if set(record.source_ids) - set(sources):
            raise ValueError(
                f"Arabic sources are unresolved for {record.spec.record_id}"
            )
        for _, target in record.spec.relations:
            if target not in record_ids:
                raise ValueError(f"Arabic graph target is unresolved: {target}")


def validate_output(text: str, records: list[DraftRecord], sources: dict[str, Any]) -> None:
    forbidden = ("WORKING DRAFT v0.1", "Prepared by", "Document Control", "Expert Review Dashboard", "RAG Test Questions", "Formal Benchmark Set", "Final Completion Report", "ESDU CONFIRMED", "\ufffd")
    for value in forbidden:
        if value in text:
            raise ValueError(f"Forbidden scaffold leaked into output: {value}")
    if text.count("\n## ") != len(records):
        raise ValueError("Every atomic record must use exactly one level-two heading")
    item_ids = {record.spec.record_id for record in records}
    validate_ontology(item_ids)
    for record in records:
        if not record.translation or any(not ARABIC_RE.search(value) for value in record.translation.values()):
            raise ValueError(f"Arabic machine draft missing for {record.spec.record_id}")
        missing = sorted(set(record.source_ids) - set(sources))
        if missing:
            raise ValueError(f"Unresolved sources for {record.spec.record_id}: {missing}")
        for _, target in record.spec.relations:
            if target not in item_ids:
                raise ValueError(f"Unresolved graph target for {record.spec.record_id}: {target}")
    if "\n# Non-retrievable source and validation appendix" not in text:
        raise ValueError("Non-retrievable appendix is missing")


def run(args: argparse.Namespace) -> None:
    checksum = hashlib.sha256(args.input.read_bytes()).hexdigest().upper()
    if checksum != SOURCE_SHA256:
        raise ValueError(f"Source DOCX checksum mismatch: {checksum}")
    chapters, chapter_sources = extract_chapters(args.input)
    records = build_records(chapters, chapter_sources, args.guide)
    apply_local_arabic_drafts(records, args.guide)
    sources = reconcile_sources(args.input, args.sources)
    text = render_markdown(records, sources, args.input.name)
    validate_output(text, records, sources)
    arabic_text = render_arabic_companion(
        records,
        sources,
        args.input.name,
        args.output.name,
    )
    validate_arabic_companion(arabic_text, records, sources)
    args.output.parent.mkdir(parents=True, exist_ok=True)
    args.output.write_text(text, encoding="utf-8", newline="\n")
    args.arabic_output.write_text(
        arabic_text,
        encoding="utf-8",
        newline="\n",
    )
    if hashlib.sha256(args.input.read_bytes()).hexdigest().upper() != SOURCE_SHA256:
        raise RuntimeError("Source DOCX changed during conversion")
    print(json.dumps({
        "output": str(args.output),
        "arabic_output": str(args.arabic_output),
        "records": len(records),
        "sources": len(sources),
        "source_sha256": checksum,
        "translation_method": "local_repository_ai_draft",
        "external_translation_calls": 0,
    }, ensure_ascii=False))


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--input", type=Path, default=Path("knowledge_base/ESDU_Agrifood_Knowledge_Base_v0.1.docx"))
    parser.add_argument("--output", type=Path, default=Path("knowledge_base/agrifood_knowledge_draft_v0.2.md"))
    parser.add_argument("--arabic-output", type=Path, default=Path("knowledge_base/agrifood_knowledge_draft_v0.2_ar.md"))
    parser.add_argument("--guide", type=Path, default=Path("knowledge_base/guide.json"))
    parser.add_argument("--sources", type=Path, default=Path("knowledge_base/sources.json"))
    args = parser.parse_args()
    try:
        run(args)
    except Exception as exc:
        print(f"conversion failed: {type(exc).__name__}: {exc}", file=sys.stderr)
        raise SystemExit(1) from exc


if __name__ == "__main__":
    main()
